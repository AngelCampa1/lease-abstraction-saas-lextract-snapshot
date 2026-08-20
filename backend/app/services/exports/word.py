"""Word document exporter using python-docx.

Generates a professional .docx lease abstraction report with cover page,
executive summary, category sections with field/value/confidence tables,
and a red flag summary appendix.
"""

from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.services.brand import BRAND_ASSETS
from app.services.exports.base import EXPORT_DISCLAIMER, ExportBase
from app.services.exports.templates import (
    DEFAULT_TEMPLATE,
    FIELD_LABELS,
    TEMPLATES,
    TemplateSection,
)
from app.services.exports.utils import (
    field_confidence_scores,
    format_value,
    get_confidence_tier,
    unwrap_field,
)

# Confidence tier colors
CONFIDENCE_COLORS: dict[str, RGBColor] = {
    "HIGH": RGBColor(0, 128, 0),
    "MEDIUM": RGBColor(255, 165, 0),
    "LOW": RGBColor(255, 0, 0),
    "N/A": RGBColor(136, 136, 136),
}

# Header row background color (light gray)
HEADER_BG_HEX = "D9E2F3"

# Backward-compatible aliases for private helper names used in existing code
_unwrap_field = unwrap_field
_get_confidence_tier = get_confidence_tier
_format_value = format_value


def _set_cell_shading(cell: Any, color_hex: str) -> None:
    """Set the background shading of a table cell.

    Args:
        cell: A python-docx table cell object.
        color_hex: Hex color string without '#' prefix.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._tc.get_or_add_tcPr(),  # noqa: SLF001
        qn("w:shd"),
    )
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


class WordExporter(ExportBase):
    """Generate Word (.docx) export of extraction results.

    Produces a structured lease abstraction report with:
    - Cover page with key identifiers
    - Executive summary with confidence overview
    - Category sections as field/value/confidence tables
    - Red flag summary appendix (if any flags detected)
    """

    @property
    def content_type(self) -> str:
        """MIME type for Word documents."""
        return (
            "application/vnd.openxmlformats-officedocument" ".wordprocessingml.document"
        )

    @property
    def extension(self) -> str:
        """File extension for Word documents."""
        return "docx"

    def generate(
        self,
        extraction_data: dict[str, Any],
        confidence_scores: dict[str, Any],
        red_flags: list[dict[str, Any]],
        template: str,
        document_filename: str,
    ) -> bytes:
        """Generate a Word document export.

        Args:
            extraction_data: Extracted lease fields keyed by field_name.
            confidence_scores: Confidence tier per field (HIGH/MEDIUM/LOW).
            red_flags: List of detected red flag dicts.
            template: Template name (e.g. 'commercial', 'office').
            document_filename: Original uploaded PDF filename.

        Returns:
            Raw bytes of the generated .docx file.
        """
        doc = Document()
        self._configure_styles(doc)
        self._add_cover_page(doc, extraction_data, document_filename)
        self._add_executive_summary(doc, extraction_data, confidence_scores, red_flags)

        template_config = TEMPLATES.get(template, TEMPLATES[DEFAULT_TEMPLATE])
        for section in template_config.sections:
            self._add_category_section(doc, section, extraction_data, confidence_scores)

        if red_flags:
            self._add_red_flag_summary(doc, red_flags)

        self._add_disclaimer(doc)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_disclaimer(self, doc: Any) -> None:
        """Add the AI-accuracy disclaimer at the end of the document.

        States that the report is AI-generated, may contain errors, must be
        verified against the source lease, and that Lextract is not liable for
        mistakes or decisions made from it.

        Args:
            doc: The python-docx Document instance.
        """
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(18)
        run = para.add_run(EXPORT_DISCLAIMER)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(89, 89, 89)

    def _configure_styles(self, doc: Any) -> None:
        """Configure document-level font styles.

        Args:
            doc: The python-docx Document instance.
        """
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

    def _add_cover_page(
        self,
        doc: Any,
        extraction_data: dict[str, Any],
        document_filename: str,
    ) -> None:
        """Add a cover page with title and key lease identifiers.

        Args:
            doc: The python-docx Document instance.
            extraction_data: Extracted lease fields.
            document_filename: Original PDF filename.
        """
        doc.add_paragraph("")

        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.add_run().add_picture(str(BRAND_ASSETS.logo_path), width=Inches(2.2))

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Lease Abstraction Report")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(31, 56, 100)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(document_filename)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(89, 89, 89)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            f"Generated: {datetime.now(UTC).strftime('%B %d, %Y')}"
        )
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(89, 89, 89)

        doc.add_paragraph("")

        key_fields = [
            ("Landlord", _unwrap_field(extraction_data.get("landlord_legal_name"))),
            ("Tenant", _unwrap_field(extraction_data.get("tenant_legal_name"))),
            ("Premises", _unwrap_field(extraction_data.get("premises_address"))),
        ]

        for label, value in key_fields:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = para.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(12)
            value_run = para.add_run(_format_value(value))
            value_run.font.size = Pt(12)

        doc.add_page_break()

    def _add_executive_summary(
        self,
        doc: Any,
        extraction_data: dict[str, Any],
        confidence_scores: dict[str, Any],
        red_flags: list[dict[str, Any]],
    ) -> None:
        """Add an executive summary section.

        Includes a brief overview of key terms, field counts by confidence
        tier, and the number of red flags detected.

        Args:
            doc: The python-docx Document instance.
            extraction_data: Extracted lease fields.
            confidence_scores: Confidence tier per field.
            red_flags: List of detected red flags.
        """
        heading = doc.add_heading("Executive Summary", level=1)
        heading.runs[0].font.color.rgb = RGBColor(31, 56, 100)

        landlord = _format_value(
            _unwrap_field(extraction_data.get("landlord_legal_name"))
        )
        tenant = _format_value(_unwrap_field(extraction_data.get("tenant_legal_name")))
        premises = _format_value(_unwrap_field(extraction_data.get("premises_address")))
        term = _format_value(_unwrap_field(extraction_data.get("lease_term_months")))
        rent = _format_value(_unwrap_field(extraction_data.get("base_rent_annual")))

        summary_text = (
            f"This lease abstraction covers the agreement between {landlord} "
            f"(Landlord) and {tenant} (Tenant) for the premises at {premises}. "
            f"The lease term is {term} months with an annual base rent of {rent}."
        )
        doc.add_paragraph(summary_text)

        total_fields = len(extraction_data)
        field_scores = field_confidence_scores(confidence_scores)
        high_count = sum(
            1 for v in field_scores.values() if _get_confidence_tier(v) == "HIGH"
        )
        medium_count = sum(
            1 for v in field_scores.values() if _get_confidence_tier(v) == "MEDIUM"
        )
        low_count = sum(
            1 for v in field_scores.values() if _get_confidence_tier(v) == "LOW"
        )
        na_count = sum(
            1 for v in field_scores.values() if _get_confidence_tier(v) == "N/A"
        )

        stats_para = doc.add_paragraph()
        stats_para.add_run("Extraction Statistics: ").bold = True
        stats_para.add_run(
            f"{total_fields} fields extracted | "
            f"{high_count} HIGH confidence | "
            f"{medium_count} MEDIUM confidence | "
            f"{low_count} LOW confidence | "
            f"{na_count} N/A (not in lease) | "
            f"{len(red_flags)} red flag(s) detected"
        )

        doc.add_paragraph("")

    def _add_category_section(
        self,
        doc: Any,
        section: TemplateSection,
        extraction_data: dict[str, Any],
        confidence_scores: dict[str, Any],
    ) -> None:
        """Add a category section with a field/value/confidence table.

        Args:
            doc: The python-docx Document instance.
            section: Template section definition.
            extraction_data: Extracted lease fields.
            confidence_scores: Confidence tier per field.
        """
        heading_level = 1 if section.emphasis else 2
        heading = doc.add_heading(section.display_name, level=heading_level)
        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(31, 56, 100)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row
        header_cells = table.rows[0].cells
        headers = ["Field", "Value", "Confidence"]
        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_shading(cell, "1F3864")

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.2)
            row.cells[1].width = Inches(3.3)
            row.cells[2].width = Inches(1.0)

        # Data rows
        for field_name in section.fields:
            raw_value = extraction_data.get(field_name)
            value = _unwrap_field(raw_value)
            confidence = _get_confidence_tier(confidence_scores.get(field_name))
            label = FIELD_LABELS.get(field_name, field_name)

            row = table.add_row()
            # Field label
            label_cell = row.cells[0]
            label_run = label_cell.paragraphs[0].add_run(label)
            label_run.font.size = Pt(9)
            label_run.bold = True

            # Value
            value_cell = row.cells[1]
            value_run = value_cell.paragraphs[0].add_run(_format_value(value))
            value_run.font.size = Pt(9)

            # Confidence with color
            conf_cell = row.cells[2]
            if confidence:
                conf_run = conf_cell.paragraphs[0].add_run(confidence)
                conf_run.font.size = Pt(9)
                conf_run.bold = True
                color = CONFIDENCE_COLORS.get(
                    str(confidence).upper(), RGBColor(128, 128, 128)
                )
                conf_run.font.color.rgb = color

        doc.add_paragraph("")

    def _add_red_flag_summary(
        self,
        doc: Any,
        red_flags: list[dict[str, Any]],
    ) -> None:
        """Add a red flag summary section at the end of the document.

        Args:
            doc: The python-docx Document instance.
            red_flags: List of red flag dicts with 'name', 'severity',
                'description' keys.
        """
        doc.add_page_break()
        heading = doc.add_heading("Appendix: Red Flags", level=1)
        heading.runs[0].font.color.rgb = RGBColor(192, 0, 0)

        doc.add_paragraph(
            f"{len(red_flags)} potential issue(s) were identified "
            "during the extraction process:"
        )

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        headers = ["Field", "Severity", "Description"]
        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_shading(cell, "8B0000")

        for flag in red_flags:
            row = table.add_row()

            field_name = flag.get("name", "Unknown")
            label = FIELD_LABELS.get(field_name, field_name)
            field_run = row.cells[0].paragraphs[0].add_run(label)
            field_run.font.size = Pt(9)

            severity = flag.get("severity", "Unknown")
            sev_run = row.cells[1].paragraphs[0].add_run(str(severity))
            sev_run.font.size = Pt(9)
            sev_run.bold = True

            message = flag.get("description", "")
            msg_run = row.cells[2].paragraphs[0].add_run(str(message))
            msg_run.font.size = Pt(9)
