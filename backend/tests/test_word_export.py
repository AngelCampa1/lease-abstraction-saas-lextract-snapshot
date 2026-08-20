"""Tests for WordExporter and export template configurations.

Exercises the Word document generation pipeline including cover page,
executive summary, category sections, red flag summary, confidence
coloring, and template validation.
"""

from io import BytesIO

import pytest
from docx import Document

from app.services.exports.base import EXPORT_DISCLAIMER, ExportBase
from app.services.exports.templates import (
    ALL_CATEGORIES,
    DEFAULT_TEMPLATE,
    FIELD_LABELS,
    TEMPLATES,
    ExportTemplate,
    TemplateSection,
)
from app.services.exports.word import (
    CONFIDENCE_COLORS,
    WordExporter,
    _format_value,
)

# -- Fixtures --


@pytest.fixture
def exporter():
    return WordExporter()


@pytest.fixture
def sample_extraction_data():
    return {
        "landlord_legal_name": "ABC Realty Corp",
        "tenant_legal_name": "Acme Inc.",
        "premises_address": "123 Main St, Suite 100, Springfield, IL 62701",
        "suite_or_unit_number": "100",
        "rentable_square_footage": 5000,
        "lease_term_months": 60,
        "base_rent_annual": "$120,000",
        "commencement_date": "2025-01-01",
        "expiration_date": "2029-12-31",
        "lease_structure_type": "NNN",
        "has_renewal_option": True,
        "cam_exclusions": ["capital improvements", "legal fees"],
        "audit_rights": True,
        "property_use_type": "Office",
    }


@pytest.fixture
def sample_confidence_scores():
    return {
        "landlord_legal_name": "HIGH",
        "tenant_legal_name": "HIGH",
        "premises_address": "HIGH",
        "suite_or_unit_number": "MEDIUM",
        "rentable_square_footage": "HIGH",
        "lease_term_months": "HIGH",
        "base_rent_annual": "HIGH",
        "commencement_date": "HIGH",
        "expiration_date": "HIGH",
        "lease_structure_type": "MEDIUM",
        "has_renewal_option": "LOW",
        "cam_exclusions": "MEDIUM",
        "audit_rights": "HIGH",
        "property_use_type": "HIGH",
    }


@pytest.fixture
def sample_red_flags():
    return [
        {
            "rule_id": "no_cam_cap",
            "name": "cam_cap_percentage",
            "severity": "HIGH",
            "description": "No CAM cap found — tenant exposed to unlimited expense increases",
            "triggered_value": "",
        },
        {
            "rule_id": "high_holdover",
            "name": "holdover_rate",
            "severity": "MEDIUM",
            "description": "Holdover rate at 200% is above market average",
            "triggered_value": "200%",
        },
    ]


def _open_docx(doc_bytes):
    """Parse docx bytes into a Document object."""
    return Document(BytesIO(doc_bytes))


# -- WordExporter property tests --


def test_word_exporter_is_export_base_subclass(exporter):
    assert isinstance(exporter, ExportBase)


def test_word_exporter_content_type(exporter):
    assert exporter.content_type == (
        "application/vnd.openxmlformats-officedocument" ".wordprocessingml.document"
    )


def test_word_exporter_extension(exporter):
    assert exporter.extension == "docx"


# -- Document generation tests --


def test_word_exporter_generates_bytes(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test_lease.pdf",
    )
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_word_exporter_generates_valid_docx(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test_lease.pdf",
    )
    doc = _open_docx(result)
    assert len(doc.paragraphs) > 0


def test_word_exporter_has_cover_page(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="my_lease.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Lease Abstraction Report" in all_text
    assert "my_lease.pdf" in all_text
    assert "ABC Realty Corp" in all_text
    assert "Acme Inc." in all_text


def test_word_exporter_has_executive_summary(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in all_text
    assert "Extraction Statistics" in all_text
    assert "HIGH confidence" in all_text
    assert "red flag(s) detected" in all_text


def test_word_exporter_has_category_sections(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # All 14 categories should appear as headings
    for section in TEMPLATES["commercial"].sections:
        assert (
            section.display_name in all_text
        ), f"Section '{section.display_name}' not found in document"


def test_word_exporter_has_red_flag_summary(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Red Flags" in all_text
    assert "2 potential issue(s)" in all_text


def test_word_exporter_labels_red_flags_as_appendix(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Appendix: Red Flags" in all_text


def test_word_exporter_no_red_flags_section_when_empty(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Red Flags" not in all_text


def test_word_exporter_includes_ai_disclaimer(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert EXPORT_DISCLAIMER in all_text
    assert "made by Lextract using AI" in all_text
    assert "Check every field against the original lease" in all_text
    assert "Lextract is not responsible for mistakes" in all_text


def test_word_exporter_includes_ai_disclaimer_without_red_flags(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert EXPORT_DISCLAIMER in all_text


def test_word_exporter_commercial_template(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    # Verify first category sections appear in commercial ordering
    commercial_order = [s.display_name for s in TEMPLATES["commercial"].sections]
    heading_order = [h for h in headings if h in commercial_order]
    assert heading_order == commercial_order


def test_word_exporter_office_template_ordering(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="office",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    office_order = [s.display_name for s in TEMPLATES["office"].sections]
    heading_order = [h for h in headings if h in office_order]
    assert heading_order == office_order


def test_word_exporter_empty_data(exporter):
    result = exporter.generate(
        extraction_data={},
        confidence_scores={},
        red_flags=[],
        template="commercial",
        document_filename="empty.pdf",
    )
    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Lease Abstraction Report" in all_text
    assert "Not found" in all_text


def test_word_exporter_confidence_colors_in_tables(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    # Find all table cell text to verify confidence labels appear
    conf_texts = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in ("HIGH", "MEDIUM", "LOW"):
                    conf_texts.add(text)
    assert "HIGH" in conf_texts
    assert "MEDIUM" in conf_texts
    assert "LOW" in conf_texts


def test_word_exporter_tables_have_field_values(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_cell_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_cell_text.append(cell.text.strip())
    assert "ABC Realty Corp" in all_cell_text
    assert "Acme Inc." in all_cell_text
    assert "5000" in all_cell_text


def test_word_exporter_unknown_template_falls_back_to_default(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="nonexistent_template",
        document_filename="test.pdf",
    )
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_word_exporter_red_flag_table_content(
    exporter, sample_extraction_data, sample_confidence_scores, sample_red_flags
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=sample_red_flags,
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_cell_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_cell_text.append(cell.text.strip())
    assert "No CAM cap found" in " ".join(all_cell_text)


# -- Template configuration tests --


def test_template_configs_have_all_categories():
    for name, tmpl in TEMPLATES.items():
        cats_in_template = {s.category for s in tmpl.sections}
        assert cats_in_template == ALL_CATEGORIES, (
            f"Template '{name}' is missing categories: "
            f"{ALL_CATEGORIES - cats_in_template}"
        )


def test_template_configs_have_14_sections():
    for name, tmpl in TEMPLATES.items():
        assert (
            len(tmpl.sections) == 14
        ), f"Template '{name}' has {len(tmpl.sections)} sections, expected 14"


def test_template_configs_no_duplicate_categories():
    for name, tmpl in TEMPLATES.items():
        cats = [s.category for s in tmpl.sections]
        assert len(cats) == len(
            set(cats)
        ), f"Template '{name}' has duplicate categories"


def test_all_templates_exist():
    assert set(TEMPLATES.keys()) == {"commercial", "office", "industrial", "retail"}


def test_default_template_exists():
    assert DEFAULT_TEMPLATE in TEMPLATES


def test_field_labels_cover_all_template_fields():
    all_fields = set()
    for tmpl in TEMPLATES.values():
        for section in tmpl.sections:
            for field_name in section.fields:
                all_fields.add(field_name)
    missing = all_fields - set(FIELD_LABELS.keys())
    assert not missing, f"Missing labels for fields: {missing}"


def test_template_section_frozen():
    section = TemplateSection(
        category="Test",
        display_name="Test",
        fields=("a", "b"),
    )
    with pytest.raises(AttributeError):
        section.category = "Changed"


def test_export_template_frozen():
    tmpl = ExportTemplate(name="test", display_name="Test")
    with pytest.raises(AttributeError):
        tmpl.name = "changed"


# -- _format_value tests --


def test_format_value_none():
    assert _format_value(None) == "Not found"


def test_format_value_bool_true():
    assert _format_value(True) == "Yes"


def test_format_value_bool_false():
    assert _format_value(False) == "No"


def test_format_value_list():
    assert _format_value(["Taxes", "Insurance", "CAM"]) == "Taxes; Insurance; CAM"


def test_format_value_empty_list():
    assert _format_value([]) == "None specified"


def test_format_value_string():
    assert _format_value("Acme Holdings LLC") == "Acme Holdings LLC"


def test_format_value_number():
    assert _format_value(42) == "42"


# -- Confidence color tests --


def test_confidence_colors_defined():
    assert "HIGH" in CONFIDENCE_COLORS
    assert "MEDIUM" in CONFIDENCE_COLORS
    assert "LOW" in CONFIDENCE_COLORS


def test_confidence_color_values():
    high = CONFIDENCE_COLORS["HIGH"]
    assert high == (0, 128, 0)
    medium = CONFIDENCE_COLORS["MEDIUM"]
    assert medium == (255, 165, 0)
    low = CONFIDENCE_COLORS["LOW"]
    assert low == (255, 0, 0)


# -- Industrial and retail template tests --


def test_word_exporter_industrial_template(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="industrial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    industrial_order = [s.display_name for s in TEMPLATES["industrial"].sections]
    heading_order = [h for h in headings if h in industrial_order]
    assert heading_order == industrial_order


def test_word_exporter_retail_template(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="retail",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    retail_order = [s.display_name for s in TEMPLATES["retail"].sections]
    heading_order = [h for h in headings if h in retail_order]
    assert heading_order == retail_order


def test_word_exporter_cover_page_date(
    exporter, sample_extraction_data, sample_confidence_scores
):
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Generated:" in all_text


def test_word_exporter_list_values_in_tables(
    exporter, sample_extraction_data, sample_confidence_scores
):
    """Verify list-type fields are formatted with semicolons."""
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=sample_confidence_scores,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_cell_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_cell_text.append(cell.text.strip())
    # cam_exclusions is ["capital improvements", "legal fees"]
    assert "capital improvements; legal fees" in all_cell_text


def test_word_exporter_stats_include_na_count(
    exporter, sample_extraction_data, sample_red_flags
):
    """Executive summary stats must include N/A count for not_found fields."""
    confidence_with_not_found = {
        "landlord_legal_name": "HIGH",
        "tenant_legal_name": "HIGH",
        "premises_address": {"score": 0.0, "tier": "not_found"},
        "base_rent_annual": {"score": 0.0, "tier": "not_found"},
        "lease_term_months": "HIGH",
    }
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=confidence_with_not_found,
        red_flags=[],
        template="commercial",
        document_filename="test.pdf",
    )
    doc = _open_docx(result)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "N/A (not in lease)" in all_text


def test_word_exporter_not_found_fields_colored_gray_in_category_table(
    exporter, sample_extraction_data
):
    """not_found fields must be labeled 'N/A' and colored gray in category tables."""
    confidence_with_not_found = {
        "landlord_legal_name": {"score": 0.95, "tier": "high"},
        "premises_address": {"score": 0.0, "tier": "not_found"},
    }
    result = exporter.generate(
        extraction_data=sample_extraction_data,
        confidence_scores=confidence_with_not_found,
        red_flags=[],
        template="commercial",
        document_filename="test.docx",
    )
    doc = _open_docx(result)
    # Find the N/A cell text in any table
    na_runs = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text == "N/A":
                            na_runs.append(run)
    assert na_runs, "Expected at least one 'N/A' run in category tables"
    for run in na_runs:
        # Gray color: RGBColor(136, 136, 136) hex = 888888
        assert run.font.color.rgb is not None
        assert str(run.font.color.rgb).upper() == "888888"
